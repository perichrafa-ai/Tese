# -*- coding: utf-8 -*-
"""
Gera um .docx com as Tabelas 3.4a–3.4d no padrão ABNT
a partir dos CSVs *formatados* (sep=';' e decimal=',').

Requisitos: pandas, python-docx
Instalação (se precisar):
    py -m pip install pandas python-docx
"""

import os
import argparse
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# --------- Configurações padrão (arquivos de entrada e metadados) ----------
DEFAULTS = {
    "coef_logit": "tabela_coef_logit_interacaoA_Base_VALIDADA_E_PRONTA__27A_prepared_formatado.csv",
    "ames_logit": "tabela_AMEs_logit_interacaoA_Base_VALIDADA_E_PRONTA__27A_prepared_formatado.csv",
    "coef_ols":   "tabela_coef_ols_interacaoB_Base_VALIDADA_E_PRONTA__27A_prepared_formatado.csv",
    "sens":       "tabela_27D_sensibilidade_Base_VALIDADA_E_PRONTA__27A_prepared_formatado.csv",
    "fig_logit":  "fig_27B_predprob_interacaoA_Base_VALIDADA_E_PRONTA__27A_prepared.png",
    "fig_ols":    "fig_27C_pred_interacaoB_Base_VALIDADA_E_PRONTA__27A_prepared.png",
    "saida_docx": "Tabelas_3.4_ABNT.docx",
}

TITULOS = {
    "3.4a": "Tabela 3.4a — Logit de reeleição com interação Gestão×Δ Abstenção",
    "3.4b": "Tabela 3.4b — Efeitos marginais médios (AMEs) do logit Gestão×Δ Abstenção",
    "3.4c": "Tabela 3.4c — OLS da variação da vantagem com interação Vantagem2016×Δ Competição",
    "3.4d": "Tabela 3.4d — Sensibilidades do modelo OLS (ΔNEC_log vs ΔNEC_diff; Top-5 por UF; HC1/cluster)"
}

NOTAS = {
    "3.4a": "Nota: Coeficientes e erros-padrão robustos (HC1). Variáveis moderadoras centradas (sufixo _c).",
    "3.4b": "Nota: AMEs calculados no valor médio das covariáveis (at=overall), com HC1.",
    "3.4c": "Nota: Erros-padrão robustos (HC1) e, quando disponível, cluster por UF. Moderadora: ΔNEC_log (principal) e ΔNEC_diff (robustez).",
    "3.4d": "Nota: Sensibilidades incluem ΔNEC_log/ΔNEC_diff, amostra Top-5 por UF e HC1/cluster por UF."
}

FONTE = "Fonte: elaboração própria."

# ----------------------- Funções utilitárias -----------------------
def _add_heading(doc: Document, texto: str, before=0.0, after=0.2, size=12, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before * 12)   # 1 linha ≈ 12pt
    p.paragraph_format.space_after = Pt(after * 12)
    run = p.add_run(texto)
    run.font.size = Pt(size)
    run.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Força fonte Times New Roman, se for sua norma
    run.font.name = "Times New Roman"
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return p

def _df_from_csv(path):
    # Lê CSV já formatado: separador ; e decimal ,
    df = pd.read_csv(path, sep=";", dtype=str)  # dtype=str mantém vírgula decimal do jeito que está
    # Limpa espaços
    df.columns = [c.strip() for c in df.columns]
    for c in df.columns:
        if df[c].dtype == object:
            df[c] = df[c].astype(str).str.strip()
    return df

def _add_table(doc: Document, df: pd.DataFrame):
    # Cria tabela Word com número de linhas = 1 (cabeçalho) + len(df)
    rows, cols = df.shape
    table = doc.add_table(rows=rows+1, cols=cols)
    table.style = 'Table Grid'
    # Larguras suaves (opcional): ajusta após colagem no Word se quiser
    # Preenche cabeçalho
    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr[j].text = str(col)
        for par in hdr[j].paragraphs:
            par.runs[0].font.bold = True
            par.runs[0].font.size = Pt(10)
            par.runs[0].font.name = "Times New Roman"
    # Preenche corpo
    for i in range(rows):
        row_cells = table.rows[i+1].cells
        for j, col in enumerate(df.columns):
            val = "" if pd.isna(df.iloc[i, j]) else str(df.iloc[i, j])
            row_cells[j].text = val
            for par in row_cells[j].paragraphs:
                if par.runs:
                    par.runs[0].font.size = Pt(10)
                    par.runs[0].font.name = "Times New Roman"
    return table

def _add_nota_e_fonte(doc: Document, nota: str):
    p_nota = doc.add_paragraph(nota)
    p_nota.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_nota.runs[0].font.size = Pt(9)
    p_nota.runs[0].font.name = "Times New Roman"

    p_fonte = doc.add_paragraph(FONTE)
    p_fonte.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_fonte.runs[0].font.size = Pt(9)
    p_fonte.runs[0].font.name = "Times New Roman"

def _add_figure_placeholder(doc: Document, titulo: str, caminho_img: str):
    _add_heading(doc, titulo, before=0.3, after=0.1, size=12, bold=True)
    if os.path.exists(caminho_img):
        # Insere a imagem
        doc.add_picture(caminho_img, width=Cm(14.5))  # ajuste fino se quiser
        cap = doc.add_paragraph("Fonte: elaboração própria.")
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.name = "Times New Roman"
    else:
        # Espaço reservado
        p = doc.add_paragraph(f"[Inserir figura disponível em: {caminho_img}]")
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.name = "Times New Roman"

# ----------------------- Pipeline principal -----------------------
def main():
    ap = argparse.ArgumentParser(description="Gerar .docx com Tabelas 3.4a–3.4d (ABNT) e figuras")
    ap.add_argument("--coef-logit", default=DEFAULTS["coef_logit"])
    ap.add_argument("--ames-logit", default=DEFAULTS["ames_logit"])
    ap.add_argument("--coef-ols",   default=DEFAULTS["coef_ols"])
    ap.add_argument("--sens",       default=DEFAULTS["sens"])
    ap.add_argument("--fig-logit",  default=DEFAULTS["fig_logit"])
    ap.add_argument("--fig-ols",    default=DEFAULTS["fig_ols"])
    ap.add_argument("--saida-docx", default=DEFAULTS["saida_docx"])
    args = ap.parse_args()

    doc = Document()

    # Tabela 3.4a
    _add_heading(doc, TITULOS["3.4a"], before=0.0, after=0.2)
    df_34a = _df_from_csv(args.coef_logit)
    _add_table(doc, df_34a)
    _add_nota_e_fonte(doc, NOTAS["3.4a"])

    doc.add_paragraph("")  # espaço

    # Tabela 3.4b (AMEs)
    _add_heading(doc, TITULOS["3.4b"], before=0.3, after=0.2)
    df_34b = _df_from_csv(args.ames_logit)
    _add_table(doc, df_34b)
    _add_nota_e_fonte(doc, NOTAS["3.4b"])

    doc.add_paragraph("")

    # Figura (logit)
    _add_figure_placeholder(doc,
        "Figura 3.x — Probabilidade predita de reeleição por níveis de gestão (−1dp, média, +1dp)",
        args.fig_logit
    )

    doc.add_paragraph("")

    # Tabela 3.4c (OLS)
    _add_heading(doc, TITULOS["3.4c"], before=0.5, after=0.2)
    df_34c = _df_from_csv(args.coef_ols)
    _add_table(doc, df_34c)
    _add_nota_e_fonte(doc, NOTAS["3.4c"])

    doc.add_paragraph("")

    # Figura (OLS)
    _add_figure_placeholder(doc,
        "Figura 3.y — Linhas de predição da variação da vantagem por níveis de vantagem inicial (−1dp, média, +1dp)",
        args.fig_ols
    )

    doc.add_paragraph("")

    # Tabela 3.4d (sensibilidades)
    _add_heading(doc, TITULOS["3.4d"], before=0.5, after=0.2)
    df_34d = _df_from_csv(args.sens)
    _add_table(doc, df_34d)
    _add_nota_e_fonte(doc, NOTAS["3.4d"])

    doc.save(args.saida_docx)
    print(f"Documento gerado: {args.saida_docx}")

if __name__ == "__main__":
    main()
# (opcional) instalar dependências
py -m pip install pandas python-docx

# gerar o docx com as tabelas e figuras
py gerar_tabelas_abnt.py
